"""
Генерация короткого, понятного плана эксперимента для научного руководителя.

Выход: thesis/План_эксперимента.docx
Формат: ~2 страницы, ГОСТ-оформление (поля, шрифт), таблица сессий.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

OUTPUT = Path(r"C:\Users\alexp\myProjects\study\thesis\План_эксперимента.docx")


def set_rfonts(run, name="Times New Roman"):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def setup(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    section = doc.sections[0]
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)


def add_para(doc, text="", *, align=None, bold=False, size=14, indent=True,
             space_before=0, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if not indent:
        pf.first_line_indent = Cm(0)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        set_rfonts(run)
    return p


def add_heading(doc, text):
    add_para(
        doc,
        text,
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        bold=True,
        indent=False,
        space_before=6,
        space_after=6,
    )


def add_bullets(doc, items):
    for item in items:
        add_para(
            doc,
            f"– {item}",
            align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            indent=True,
            space_after=0,
        )


def add_sessions_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header = table.rows[0].cells
    headers = ["№", "Конфигурация", "Цель", "Что измеряем"]
    for cell, title in zip(header, headers):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        set_rfonts(run)

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        values = [str(i)] + list(row)
        for cell, value in zip(cells, values):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            set_rfonts(run)


def build():
    doc = Document()
    setup(doc)

    add_para(
        doc,
        "План эксперимента",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        bold=True,
        size=16,
        indent=False,
        space_after=6,
    )
    add_para(
        doc,
        "К магистерской ВКР: устранение помеховых сигналов при "
        "измерении диаграмм обратного рассеяния в безэховых камерах "
        "с неудовлетворительными безэховыми условиями",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        indent=False,
        space_after=18,
    )

    add_heading(doc, "1. Цель эксперимента")
    add_para(
        doc,
        "Экспериментально подтвердить работоспособность предлагаемого "
        "гибридного метода (ISAR-фокусировка + VLM-атрибуция помех) при "
        "измерении ЭПР в безэховой камере с искусственно ухудшенными "
        "безэховыми условиями; сравнить метод с классическими подходами "
        "по количественным метрикам PSR, разрешение и погрешность ЭПР.",
    )

    add_heading(doc, "2. Место проведения")
    add_para(
        doc,
        "Безэховая камера ПАО «Радиофизика»; диапазон — X (8–12 ГГц); "
        "среда обработки — Python/MATLAB; VLM — Claude API / "
        "локальная Qwen2-VL.",
    )

    add_heading(doc, "3. Оборудование и мишени")
    add_bullets(
        doc,
        [
            "безэховая камера с известными характеристиками поглощения;",
            "векторный анализатор цепей или когерентный радарный приёмник "
            "с записью IQ-сигнала;",
            "рупорные антенны передачи и приёма;",
            "поворотная установка (турникет) с точностью позиционирования "
            "не хуже 0,5°;",
            "эталонные мишени: металлическая сфера (D = 100 мм), "
            "уголковый отражатель; реальный тестовый объект "
            "(согласуется с научным руководителем).",
        ],
    )

    add_heading(doc, "4. Серии измерений")
    add_sessions_table(
        doc,
        [
            (
                "Базовая БЭК, без мишени",
                "«Подпись» камеры",
                "фон, уровень помех",
            ),
            (
                "Базовая БЭК, сфера",
                "Калибровка, эталон",
                "ДОР сферы, 360°",
            ),
            (
                "Базовая БЭК, уголковый отражатель",
                "Второй эталон",
                "ДОР уголка, 360°",
            ),
            (
                "«Ухудшенная» БЭК, без мишени",
                "Фон плохих условий",
                "уровень и структура помех",
            ),
            (
                "«Ухудшенная» БЭК, сфера",
                "Основной кейс метода",
                "ДОР сферы в помехах",
            ),
            (
                "«Ухудшенная» БЭК, уголковый отражатель",
                "Второй кейс",
                "ДОР уголка в помехах",
            ),
            (
                "«Ухудшенная» БЭК, реальный объект",
                "Практическая демонстрация",
                "ДОР объекта",
            ),
        ],
    )
    add_para(
        doc,
        "«Ухудшение» камеры достигается удалением части поглотителей "
        "или добавлением металлического экрана; если условия лаборатории "
        "этого не позволяют — мишень размещается вблизи стены, пола или "
        "оснастки так, чтобы выраженно проявлялся мультипут.",
        indent=True,
        space_before=6,
    )

    add_heading(doc, "5. Параметры измерения")
    add_bullets(
        doc,
        [
            "диапазон частот: 8–12 ГГц, шаг по частоте 10 МГц;",
            "поляризация: линейная, H-H или V-V;",
            "шаг по углу наблюдения: 1° (при необходимости 0,5°);",
            "диапазон углов: полный 360° или сектор ±90° относительно оси;",
            "количество реализаций на точку: 4–8 для усреднения.",
        ],
    )

    add_heading(doc, "6. Пайплайн обработки данных")
    add_bullets(
        doc,
        [
            "ISAR-фокусировка по записанным IQ-данным (кросс-корреляция "
            "Doppler-профилей, фазовая компенсация);",
            "подавление помех: временное стробирование, SVD-разложение, "
            "адаптивный LMS/RLS-фильтр;",
            "восстановление диаграммы обратного рассеяния;",
            "VLM-атрибуция помеховых источников по range-Doppler "
            "изображению;",
            "расчёт метрик PSR, разрешение, погрешность ЭПР;",
            "генерация диагностического отчёта средствами LLM.",
        ],
    )

    add_heading(doc, "7. Критерии оценки")
    add_bullets(
        doc,
        [
            "PSR = 10·lg(Pпаразит / Pпосле_фильтрации), дБ — чем больше, "
            "тем лучше подавление;",
            "разрешающая способность: ширина главного лепестка и уровень "
            "боковых максимумов до/после обработки;",
            "точность измерения ЭПР: среднеквадратическое отклонение "
            "измеренной ДОР от теоретической на эталонных целях;",
            "скорость обработки: время полного прогона на стандартном "
            "ноутбуке.",
        ],
    )

    add_heading(doc, "8. Ориентировочные сроки")
    add_bullets(
        doc,
        [
            "неделя 1: согласование доступа, калибровка аппаратуры, "
            "тестовая сессия;",
            "недели 2–3: основной объём измерений (сессии 1–7);",
            "недели 3–4: обработка данных, сравнение методов, VLM-"
            "экспертиза;",
            "неделя 5: подготовка результатов и графиков для ВКР, "
            "интеграция в веб-приложение.",
        ],
    )

    add_heading(doc, "9. Риски и резервные варианты")
    add_bullets(
        doc,
        [
            "нет доступа к «ухудшению» БЭК — моделирование «плохих» "
            "условий расположением мишени у стен или численной "
            "симуляцией;",
            "нет штатного ISAR-пайплайна на стенде — офлайн-обработка "
            "сырых IQ на рабочем ноутбуке;",
            "ограниченное число сеансов — приоритет отдаётся сессиям "
            "№ 5 и № 6 как ключевым для демонстрации метода.",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"[ok] saved: {OUTPUT}")


if __name__ == "__main__":
    build()
