"""
Генерация шаблона магистерской диссертации по ГОСТ 7.32-2017 / Р 7.0.11-2011.

Выход: thesis/Диссертация_шаблон.docx

Что настраивается:
- поля: левое 30 мм, правое 15 мм, верхнее/нижнее 20 мм;
- основной шрифт Times New Roman 14 pt, межстрочный 1,5, абзац 1,25 см;
- структурные заголовки (ВВЕДЕНИЕ, ГЛАВА N, ЗАКЛЮЧЕНИЕ, СПИСОК ЛИТЕРАТУРЫ,
  ПРИЛОЖЕНИЯ) — прописными, по центру, полужирным, с новой страницы;
- титульный лист по образцу МФТИ с актуальной темой и научруком;
- нумерация страниц — снизу по центру, со 2-й страницы.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

OUTPUT = Path(r"C:\Users\alexp\myProjects\study\thesis\Диссертация_шаблон.docx")

TOPIC = (
    "«Устранение помеховых сигналов при измерении диаграмм обратного "
    "рассеяния в безэховых камерах с неудовлетворительными "
    "безэховыми условиями»"
)
AUTHOR = "Певненко Александр Александрович"
GROUP = "М01-402б"
SUPERVISOR = "Елизаров Сергей Валерьевич"
SUPERVISOR_DEGREE = "кандидат технических наук"
YEAR = "2026"
CITY = "Москва"


def set_rfonts(run, name="Times New Roman"):
    """Выставить гарнитуру для всех кодовых страниц (кириллица ↔ латиница)."""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def set_page_numbers(section):
    """Номер страницы внизу по центру, арабскими цифрами."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    set_rfonts(run)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE   \\* MERGEFORMAT"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def configure_title_section(section):
    """Первая секция — титул, номер не показываем."""
    section.different_first_page_header_footer = True


def set_margins(section):
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)


def set_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")

    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def add_para(
    doc,
    text="",
    *,
    align=None,
    bold=False,
    size=14,
    indent=True,
    space_before=0,
    space_after=0,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if not indent:
        pf.first_line_indent = Cm(0)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        set_rfonts(run)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_structural_heading(doc, text):
    """Заголовок структурного элемента: ВСЕ ПРОПИСНЫЕ, центр, жирный, без номера."""
    add_para(
        doc,
        text.upper(),
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        bold=True,
        indent=False,
        space_after=18,
    )


def add_chapter_heading(doc, number, text):
    """Заголовок главы: ГЛАВА N по центру, на следующей строке — название."""
    add_para(
        doc,
        f"ГЛАВА {number}",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        bold=True,
        indent=False,
        space_after=6,
    )
    add_para(
        doc,
        text.upper(),
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        bold=True,
        indent=False,
        space_after=18,
    )


def add_section_heading(doc, numbering, text):
    """Подраздел: N.M Название, слева с абзацного отступа, полужирный."""
    add_para(
        doc,
        f"{numbering} {text}",
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        bold=True,
        indent=True,
        space_before=12,
        space_after=6,
    )


def add_placeholder(doc, text):
    add_para(doc, text, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, indent=True)


def build_title_page(doc):
    add_para(
        doc,
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        bold=True,
        indent=False,
    )
    add_para(
        doc,
        "Федеральное государственное автономное образовательное "
        "учреждение высшего образования",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        indent=False,
    )
    add_para(
        doc,
        "«Московский физико-технический институт "
        "(национальный исследовательский университет)»",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        bold=True,
        indent=False,
    )
    add_para(
        doc,
        "Физтех-школа радиотехники и компьютерных технологий",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        indent=False,
    )
    add_para(
        doc,
        "Кафедра радиофизики и технической кибернетики",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        indent=False,
    )

    for _ in range(4):
        add_para(doc, indent=False)

    add_para(
        doc,
        "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА МАГИСТРА",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        bold=True,
        indent=False,
        space_after=6,
    )
    add_para(
        doc,
        "на тему:",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        indent=False,
        space_after=6,
    )
    add_para(
        doc,
        TOPIC,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        bold=True,
        size=16,
        indent=False,
        space_after=6,
    )

    for _ in range(6):
        add_para(doc, indent=False)

    # Блок «студент / научный руководитель» — таблица 2 колонки
    add_para(
        doc,
        f"Студент группы {GROUP}:",
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        indent=False,
    )
    add_para(
        doc,
        f"_______________________ {AUTHOR}",
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        indent=False,
        space_after=12,
    )
    add_para(
        doc,
        "Научный руководитель:",
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        indent=False,
    )
    add_para(
        doc,
        f"_______________________ {SUPERVISOR},",
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        indent=False,
        space_after=0,
    )
    add_para(
        doc,
        SUPERVISOR_DEGREE,
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        indent=False,
        space_after=12,
    )

    for _ in range(4):
        add_para(doc, indent=False)

    add_para(
        doc,
        f"{CITY} {YEAR}",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        indent=False,
    )


def build_content_page(doc):
    add_structural_heading(doc, "Содержание")
    toc_items = [
        "ВВЕДЕНИЕ",
        "ГЛАВА 1. АНАЛИТИЧЕСКИЙ ОБЗОР",
        "ГЛАВА 2. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ",
        "ГЛАВА 3. РАЗРАБОТКА ГИБРИДНОГО МЕТОДА И ЕГО РЕАЛИЗАЦИЯ",
        "ГЛАВА 4. ЭКСПЕРИМЕНТАЛЬНЫЕ ИССЛЕДОВАНИЯ",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ЛИТЕРАТУРЫ",
        "ПРИЛОЖЕНИЕ А. Листинги программ",
        "ПРИЛОЖЕНИЕ Б. Результаты измерений",
        "ПРИЛОЖЕНИЕ В. Техническое задание на эксперимент",
        "ПРИЛОЖЕНИЕ Г. Документация и ссылка на веб-приложение",
    ]
    for item in toc_items:
        add_para(
            doc,
            item,
            align=WD_PARAGRAPH_ALIGNMENT.LEFT,
            indent=False,
            space_after=6,
        )


def build_intro(doc):
    add_structural_heading(doc, "Введение")
    add_placeholder(
        doc,
        "Актуальность темы. [Обосновать актуальность: рост объёмов "
        "измерений ЭПР в условиях ограниченных возможностей по качеству "
        "безэховых камер; высокая стоимость идеальных БЭК; практическая "
        "потребность в методах компенсации помеховых сигналов.]",
    )
    add_placeholder(
        doc,
        "Цель работы — разработка гибридного метода обработки сигналов "
        "при измерении диаграмм обратного рассеяния в безэховых камерах "
        "с неудовлетворительными безэховыми условиями, сочетающего "
        "классическую ISAR-фокусировку с автоматизированной атрибуцией "
        "помеховых сигналов на базе визуально-языковой модели.",
    )
    add_placeholder(doc, "Задачи:")
    for task in (
        "провести обзор существующих методов подавления помех при "
        "измерении ЭПР в БЭК;",
        "разработать математическую модель сигналов и помех в БЭК с "
        "неудовлетворительными безэховыми условиями;",
        "разработать гибридный алгоритм обработки, сочетающий ISAR-"
        "фокусировку и VLM-атрибуцию помеховых источников;",
        "реализовать алгоритм в виде программного комплекса и "
        "интерактивного веб-приложения;",
        "провести экспериментальные исследования метода в безэховой "
        "камере ПАО «Радиофизика»;",
        "оценить эффективность метода по метрикам PSR, разрешающая "
        "способность и погрешность измерения ЭПР.",
    ):
        add_para(
            doc,
            f"– {task}",
            align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            indent=True,
        )
    add_placeholder(
        doc,
        "Объект исследования — электромагнитные сигналы, отражённые "
        "от стандартных радиолокационных мишеней в безэховой камере "
        "с недостаточным поглощением отражений.",
    )
    add_placeholder(
        doc,
        "Предмет исследования — процессы формирования, искажения и "
        "подавления помеховых сигналов при измерении диаграмм обратного "
        "рассеяния в условиях ухудшенной безэховой среды.",
    )
    add_placeholder(
        doc,
        "Научная новизна. [Сформулировать 3–5 положений на защиту, "
        "опираясь на документ «Новизна и положения на защиту».]",
    )
    add_placeholder(
        doc,
        "Практическая значимость. [Открытое веб-приложение, реализующее "
        "метод; методика проведения измерений ЭПР в БЭК с "
        "неудовлетворительными условиями; результаты экспериментальной "
        "апробации на базе ПАО «Радиофизика».]",
    )
    add_placeholder(
        doc,
        "Структура работы. Диссертация состоит из введения, четырёх "
        "глав, заключения, списка литературы и приложений.",
    )


def build_chapter1(doc):
    add_chapter_heading(doc, 1, "Аналитический обзор")
    sections = [
        ("1.1", "Безэховые камеры: классификация и характеристики"),
        ("1.2", "Помехи в БЭК с неудовлетворительными условиями: механизмы и классификация"),
        ("1.3", "Классические методы подавления помеховых сигналов"),
        ("1.3.1", "Временное стробирование"),
        ("1.3.2", "Частотная фильтрация и оконные функции"),
        ("1.3.3", "Фоновое вычитание и калибровка"),
        ("1.3.4", "Адаптивные фильтры (LMS, RLS)"),
        ("1.3.5", "Методы разложения (SVD, PCA, CLEAN)"),
        ("1.4", "ISAR в задачах метрологии ЭПР"),
        ("1.5", "Применение машинного обучения и языковых моделей в радиолокации"),
        ("1.6", "Постановка задачи диссертации"),
    ]
    for num, title in sections:
        add_section_heading(doc, num, title)
        add_placeholder(doc, "[Текст раздела.]")


def build_chapter2(doc):
    add_chapter_heading(doc, 2, "Теоретические основы")
    sections = [
        ("2.1", "Модель сигнала измерения ЭПР в БЭК"),
        ("2.2", "Модель помеховых сигналов: мультипут, переотражения, дифракция"),
        ("2.3", "ISAR: range-Doppler-отображение и фокусировка"),
        ("2.4", "Архитектура визуально-языковой модели и её применимость "
                 "к анализу радиолокационных изображений"),
        ("2.5", "Метрики качества: PSR, разрешение, погрешность ЭПР"),
        ("2.6", "Выводы по главе"),
    ]
    for num, title in sections:
        add_section_heading(doc, num, title)
        add_placeholder(doc, "[Текст раздела.]")


def build_chapter3(doc):
    add_chapter_heading(doc, 3, "Разработка гибридного метода и его реализация")
    sections = [
        ("3.1", "Постановка задачи и структура предлагаемого пайплайна"),
        ("3.2", "Сигнальная часть: ISAR-фокусировка и подавление помех"),
        ("3.3", "Аналитическая часть: VLM-атрибуция помеховых сигналов"),
        ("3.3.1", "Таксономия классов помеховых источников"),
        ("3.3.2", "Формат входа VLM и построение промпта"),
        ("3.3.3", "Оценка качества атрибуции"),
        ("3.4", "Формирование итогового диагностического отчёта"),
        ("3.5", "Программная реализация"),
        ("3.6", "Веб-приложение для интерактивной обработки данных"),
        ("3.7", "Верификация на численной модели"),
        ("3.8", "Выводы по главе"),
    ]
    for num, title in sections:
        add_section_heading(doc, num, title)
        add_placeholder(doc, "[Текст раздела.]")


def build_chapter4(doc):
    add_chapter_heading(doc, 4, "Экспериментальные исследования")
    sections = [
        ("4.1", "Описание экспериментальной установки"),
        ("4.2", "План и параметры измерений"),
        ("4.3", "Методика проведения и сбора данных"),
        ("4.4", "Результаты измерений: сырые IQ, РЛИ, диаграммы рассеяния"),
        ("4.5", "Сравнительный анализ методов подавления помех"),
        ("4.6", "Количественная оценка по метрикам"),
        ("4.7", "Сопоставление VLM-атрибуции с экспертной разметкой"),
        ("4.8", "Сопоставление результатов измерений и расчётов веб-приложения"),
        ("4.9", "Выводы по главе"),
    ]
    for num, title in sections:
        add_section_heading(doc, num, title)
        add_placeholder(doc, "[Текст раздела.]")


def build_conclusion(doc):
    add_structural_heading(doc, "Заключение")
    add_placeholder(doc, "[Основные результаты работы, соответствие поставленным задачам, "
                          "направления дальнейших исследований.]")


def build_references(doc):
    add_structural_heading(doc, "Список литературы")
    add_placeholder(
        doc,
        "[Список оформляется по ГОСТ Р 7.0.5-2008 и ГОСТ 7.1-2003. "
        "Источники нумеруются в порядке появления ссылок в тексте или "
        "в алфавитном порядке.]",
    )


def build_appendices(doc):
    for letter, title in [
        ("А", "Листинги программ"),
        ("Б", "Результаты измерений"),
        ("В", "Техническое задание на эксперимент"),
        ("Г", "Документация и ссылка на веб-приложение"),
    ]:
        add_structural_heading(doc, f"Приложение {letter}")
        add_para(
            doc,
            title,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            bold=True,
            indent=False,
            space_after=12,
        )
        add_placeholder(doc, "[Содержимое приложения.]")
        add_page_break(doc)


def build():
    doc = Document()
    set_normal_style(doc)

    section = doc.sections[0]
    set_margins(section)
    configure_title_section(section)
    set_page_numbers(section)

    build_title_page(doc)
    add_page_break(doc)

    build_content_page(doc)
    add_page_break(doc)

    build_intro(doc)
    add_page_break(doc)

    build_chapter1(doc)
    add_page_break(doc)
    build_chapter2(doc)
    add_page_break(doc)
    build_chapter3(doc)
    add_page_break(doc)
    build_chapter4(doc)
    add_page_break(doc)

    build_conclusion(doc)
    add_page_break(doc)

    build_references(doc)
    add_page_break(doc)

    build_appendices(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"[ok] saved: {OUTPUT}")


if __name__ == "__main__":
    build()
